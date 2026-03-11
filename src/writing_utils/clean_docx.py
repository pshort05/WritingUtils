#!/usr/bin/env python3
"""
clean_docx.py - Format and clean Word .docx files for publishing.

Features (all optional, combinable freely):
  --clean              Remove extra empty paragraphs within the range
  --page-breaks        Insert a page break before each Heading 1 within the range
  --fix-hrules         Replace fixed-width HR shapes with margin-relative paragraph borders
  --indent-paragraphs  Add a first-line indent to body text paragraphs within the range
  --font-to            Convert all non-skipped fonts to the specified font
  --font-skip          Font(s) to leave unchanged (repeatable)

A YAML config file can supply any of these options:
  -c config.yaml

Precedence (highest wins): command-line flags > config file > built-in defaults

Requirements:
    pip install python-docx pyyaml

Usage:
    python3 clean_docx.py -c thunderwing.yaml
    python3 clean_docx.py -i input.docx -o output.docx --clean
    python3 clean_docx.py -c thunderwing.yaml --font-to "Courier Prime"

Config file format (thunderwing.yaml):
    input:       "Thunderwing Book 1.docx"
    output:      "Thunderwing Book 1 - clean.docx"
    start:       "Prologue"
    end:         "The Luminarch"
    clean:       true
    page-breaks: true
    fix-hrules:  true
    font-to:     "Times New Roman"
    font-skip:
      - "Roboto Mono"
    log-level:   INFO        # NONE | ERROR | INFO | DEBUG  (default: NONE)
    log-file:    clean.log   # omit to log to stderr
"""

import argparse
import logging
import os
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from writing_utils._util import load_config, setup_logging

log = logging.getLogger(__name__)


def apply_config(args, config):
    """Fill None arg values from config.  CLI-supplied values always win."""
    scalars = {
        "input":             "input",
        "output":            "output",
        "start":             "start",
        "end":               "end",
        "clean":             "clean",
        "page_breaks":       "page-breaks",
        "fix_hrules":        "fix-hrules",
        "indent_paragraphs": "indent-paragraphs",
        "excerpt_font":      "excerpt-font",
        "font_to":           "font-to",
        "log_level":         "log-level",
        "log_file":          "log-file",
    }
    for dest, key in scalars.items():
        if getattr(args, dest) is None and key in config:
            setattr(args, dest, config[key])

    # font-skip is a list; only pull from config if CLI provided nothing
    if args.font_skip is None:
        args.font_skip = list(config.get("font-skip") or [])

    # Apply built-in defaults for anything still unset
    for dest in ("clean", "page_breaks", "fix_hrules", "indent_paragraphs"):
        if getattr(args, dest) is None:
            setattr(args, dest, False)
    if args.font_skip is None:
        args.font_skip = []

    return args


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def find_heading(paragraphs, search_text):
    """Return the index of the first heading containing search_text (case-insensitive)."""
    needle = search_text.strip().lower()
    for i, p in enumerate(paragraphs):
        if p.style.name.startswith("Heading") and needle in p.text.strip().lower():
            return i
    return None


def resolve_range(doc, args, verbose=False):
    """Return (start_idx, end_idx) based on --start / --end args."""
    paras = doc.paragraphs
    total = len(paras)
    start_idx = 0
    end_idx = total - 1

    if args.start:
        idx = find_heading(paras, args.start)
        if idx is None:
            msg = f"No heading found containing '{args.start}'"
            log.error(msg)
            print(f"Error: {msg}", file=sys.stderr)
            sys.exit(1)
        start_idx = idx
        if verbose:
            print(f"Start : para {start_idx:5}  \"{paras[start_idx].text.strip()[:70]}\"")
        log.info("Range start: para %d  \"%s\"", start_idx, paras[start_idx].text.strip()[:70])

    if args.end:
        idx = find_heading(paras, args.end)
        if idx is None:
            msg = f"No heading found containing '{args.end}'"
            log.error(msg)
            print(f"Error: {msg}", file=sys.stderr)
            sys.exit(1)
        end_idx = idx - 1
        if verbose:
            print(f"End   : para {end_idx:5}  (just before \"{paras[idx].text.strip()[:70]}\")")
        log.info("Range end: para %d  (just before \"%s\")", end_idx, paras[idx].text.strip()[:70])

    if verbose:
        print(f"Range : {end_idx - start_idx + 1} paragraphs  (doc total: {total})")
    log.info("Range: %d paragraphs  (doc total: %d)", end_idx - start_idx + 1, total)
    return start_idx, end_idx


# ---------------------------------------------------------------------------
# Feature: clean empty paragraphs
# ---------------------------------------------------------------------------

def _has_bottom_border(p):
    """Return True if the paragraph has a non-nil bottom border (i.e. is an HR)."""
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        return False
    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        return False
    return bottom.get(qn("w:val"), "nil") not in ("nil", "none", "")


def collect_removals(paragraphs, start_idx, end_idx):
    """Walk the paragraph range and return a list of XML elements to remove.
    Paragraphs with a bottom border (horizontal rules) are never removed."""
    to_remove = []
    total = end_idx - start_idx + 1
    report_interval = max(1, total // 10)
    i = start_idx

    while i <= end_idx:
        if (i - start_idx) % report_interval == 0:
            pct = (i - start_idx) * 100 // total
            print(f"  Scanning paragraphs... {pct}% ({i - start_idx}/{total})", end="\r")

        p = paragraphs[i]
        is_empty = not p.text.strip()
        style_name = p.style.name
        is_heading = style_name.startswith("Heading")

        # Never remove a paragraph that carries a horizontal rule border
        if is_empty and _has_bottom_border(p):
            log.debug("Para %d: keeping HR border paragraph", i)
            i += 1
            continue

        # Rule 1: empty Heading 1 paragraph — remove (export artifact).
        # Empty Heading 2+ paragraphs are kept; they may be visual
        # scene-break separators styled by the heading definition.
        if is_empty and style_name == "Heading 1":
            log.debug("Para %d: removing empty Heading 1", i)
            to_remove.append(p._element)
            i += 1
            continue

        # Empty Heading 2+ — preserve and advance (not an artifact)
        if is_empty and is_heading:
            log.debug("Para %d: keeping empty %s (scene-break separator)", i, style_name)
            i += 1
            continue

        # Rules 2 & 3: run of empty normal paragraphs
        if is_empty:
            j = i
            while (
                j <= end_idx
                and not paragraphs[j].text.strip()
                and not paragraphs[j].style.name.startswith("Heading")
            ):
                j += 1

            run_len = j - i
            if run_len == 1:
                log.debug("Para %d: removing single empty paragraph", i)
                to_remove.append(paragraphs[i]._element)
            else:
                log.debug("Para %d-%d: scene break — keeping first, removing %d", i, j - 1, run_len - 1)
                for k in range(i + 1, j):
                    to_remove.append(paragraphs[k]._element)

            i = j
            continue

        i += 1

    print(f"  Scanning paragraphs... done ({total}/{total})         ")
    return to_remove


# ---------------------------------------------------------------------------
# Feature: fix horizontal rules
# ---------------------------------------------------------------------------

_AC_TAG = "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent"


def _has_page_break_run(elem):
    """Return True if an XML element contains a <w:br w:type="page"/> run."""
    for r in elem.findall(qn("w:r")):
        for br in r.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
    return False


def _apply_hr_border(p):
    """Add a margin-relative bottom border to paragraph p."""
    pPr = p._p.get_or_add_pPr()
    old_pBdr = pPr.find(qn("w:pBdr"))
    if old_pBdr is not None:
        pPr.remove(old_pBdr)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")       # 0.75pt — matches original visual weight
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "A0A0A0")
    pBdr.append(bottom)
    rPr_elem = pPr.find(qn("w:rPr"))
    if rPr_elem is not None:
        rPr_elem.addprevious(pBdr)
    else:
        pPr.append(pBdr)
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        rPr_elem2 = pPr.find(qn("w:rPr"))
        if rPr_elem2 is not None:
            rPr_elem2.addprevious(ind)
        else:
            pPr.append(ind)
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")


def fix_hrules(paragraphs, start_idx, end_idx):
    """Replace fixed-width horizontal line drawing shapes with paragraph bottom
    borders, which automatically stretch to whatever the current text margins are.

    Also adds a bottom border to empty Heading 2+ paragraphs that have no border.
    These are used as scene-break separators; without an explicit border they
    render as invisible whitespace rather than a visible rule.

    Google Docs exports '---' as an inline drawing shape (mc:AlternateContent)
    with a hardcoded pixel width. This breaks when margins change for 2-sided
    printing. A w:pBdr bottom border is margin-relative and always correct.

    Returns the number of horizontal lines fixed.
    """
    count = 0
    for i in range(start_idx, end_idx + 1):
        p = paragraphs[i]

        # Case 1: AlternateContent HR drawing shape (Google Docs '---' export)
        hr_run = None
        for r in p._p.findall(qn("w:r")):
            if r.find(_AC_TAG) is not None:
                hr_run = r
                break
        if hr_run is not None:
            p._p.remove(hr_run)
            _apply_hr_border(p)
            log.debug("Para %d: converted HR drawing shape to paragraph border", i)
            count += 1
            continue

        # Case 2: empty Heading 2+ paragraph used as a scene-break separator.
        # The Heading 2 style carries no border, so without an explicit pBdr
        # these paragraphs render as blank space rather than a visible rule.
        # Skip paragraphs that already contain a page-break run — those are
        # chapter-break containers, not scene separators.
        if (not p.text.strip()
                and p.style.name.startswith("Heading")
                and p.style.name != "Heading 1"
                and not _has_bottom_border(p)
                and not _has_page_break_run(p._p)):
            _apply_hr_border(p)
            log.debug("Para %d: added border to empty %s scene-break paragraph", i, p.style.name)
            count += 1

    return count


# ---------------------------------------------------------------------------
# Feature: page breaks
# ---------------------------------------------------------------------------

def insert_page_breaks(paragraphs, start_idx, end_idx):
    """Insert an explicit page break paragraph immediately before each
    non-empty Heading 1 within the range. Skips if the immediately preceding
    sibling already contains a page-break run. Returns the count inserted."""
    count = 0
    for i in range(start_idx, end_idx + 1):
        p = paragraphs[i]
        if p.style.name == "Heading 1" and p.text.strip():
            prev = p._element.getprevious()
            if prev is not None and _has_page_break_run(prev):
                log.debug("Para %d: skipping page break insert (already present before \"%s\")",
                          i, p.text.strip()[:60])
                continue
            new_p = OxmlElement("w:p")
            new_r = OxmlElement("w:r")
            new_br = OxmlElement("w:br")
            new_br.set(qn("w:type"), "page")
            new_r.append(new_br)
            new_p.append(new_r)
            p._element.addprevious(new_p)
            log.debug("Para %d: inserted page break before \"%s\"", i, p.text.strip()[:60])
            count += 1
    return count


# ---------------------------------------------------------------------------
# Feature: paragraph indentation
# ---------------------------------------------------------------------------

_SKIP_ALIGNMENTS = {
    WD_ALIGN_PARAGRAPH.CENTER,
    WD_ALIGN_PARAGRAPH.RIGHT,
    WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}

# ~3 characters wide at 12pt body text
_FIRST_LINE_INDENT = Inches(0.25)


def _is_title_component(p):
    """Return True if every non-empty run in the paragraph is bold, italic, or
    both. Identifies subtitle paragraphs (place names, character names, dates,
    POV labels) that follow a chapter heading but use the normal paragraph style.
    Handles paragraphs that mix bold and italic runs (e.g. bold location + italic
    character name in the same paragraph)."""
    runs = [r for r in p.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.bold or r.italic for r in runs)


def indent_paragraphs(paragraphs, start_idx, end_idx, excerpt_font=None):
    """Add a first-line indent to body text paragraphs within the range.

    Skips:
      - Empty paragraphs
      - Heading paragraphs (any Heading style)
      - Center-justified, right-justified, and distributed paragraphs
      - Title-component paragraphs: every non-empty run is bold, italic, or
        both, appearing immediately after a heading (chapter subtitles such as
        place names, character names, POV labels)
      - The first true body paragraph after each heading block, per publishing
        convention (opening paragraph of a chapter or scene is flush left)
      - Paragraphs using excerpt_font (block-quoted excerpts get their own
        indentation from format_excerpts and should not get first-line indent)

    Returns the number of paragraphs indented.
    """
    count = 0
    after_heading = True  # start in heading-block state to catch chapter subtitles

    for i in range(start_idx, end_idx + 1):
        p = paragraphs[i]

        if not p.text.strip():
            # An empty heading paragraph signals a scene break; re-enter
            # heading-block state so the first body paragraph after it is
            # also left un-indented (publishing convention).
            if p.style.name.startswith("Heading"):
                after_heading = True
            continue

        if p.style.name.startswith("Heading"):
            after_heading = True
            log.debug("Para %d: skipping heading for indent", i)
            continue

        if p.paragraph_format.alignment in _SKIP_ALIGNMENTS:
            log.debug("Para %d: skipping non-left-aligned paragraph for indent", i)
            continue

        if excerpt_font and _paragraph_uses_font(p, excerpt_font):
            log.debug("Para %d: skipping excerpt paragraph for indent", i)
            continue

        if after_heading:
            if _is_title_component(p):
                log.debug("Para %d: skipping title-component paragraph for indent", i)
                continue  # remain in after_heading state
            # First true body paragraph after heading block — no indent
            # (publishing convention: opening paragraph of chapter/scene
            # is flush left).
            after_heading = False
            log.debug("Para %d: skipping first paragraph after heading (no indent)", i)
            continue

        p.paragraph_format.first_line_indent = _FIRST_LINE_INDENT
        log.debug("Para %d: applied first-line indent", i)
        count += 1

    return count


# ---------------------------------------------------------------------------
# Feature: excerpt block formatting
# ---------------------------------------------------------------------------

_EXCERPT_LEFT_INDENT  = Inches(0.5)
_EXCERPT_RIGHT_INDENT = Inches(0.5)


def _paragraph_uses_font(p, font_name):
    """Return True if any w:rFonts element in the paragraph references font_name.

    Checks both run-level and paragraph-level (pPr/rPr) font settings.
    """
    font_lower = font_name.lower()
    font_attrs = (qn("w:ascii"), qn("w:hAnsi"), qn("w:cs"), qn("w:eastAsia"))
    for rFonts in p._p.findall(".//" + qn("w:rFonts")):
        for attr in font_attrs:
            val = rFonts.get(attr, "")
            if val.lower() == font_lower:
                return True
    return False


def _elem_is_empty_para(elem):
    """Return True if elem is a <w:p> element containing no visible text."""
    if elem.tag != qn("w:p"):
        return False
    for t in elem.findall(".//" + qn("w:t")):
        if t.text and t.text.strip():
            return False
    return True


def format_excerpts(paragraphs, start_idx, end_idx, excerpt_font):
    """Apply block-quote formatting to embedded written excerpts.

    Excerpts are identified by font name.  Each contiguous group of excerpt
    paragraphs (empty gaps between them are included) receives:
      - 0.5" left indent and 0.5" right indent
      - An empty paragraph inserted before the group if one is not already present
      - An empty paragraph inserted after the group if one is not already present

    Returns the number of excerpt paragraphs formatted.
    """
    # Identify every excerpt paragraph in range
    excerpt_set = set()
    for i in range(start_idx, end_idx + 1):
        if _paragraph_uses_font(paragraphs[i], excerpt_font):
            excerpt_set.add(i)

    if not excerpt_set:
        return 0

    excerpt_indices = sorted(excerpt_set)

    # Build groups: consecutive excerpt indices where gaps contain only empty paras
    groups = []
    group_start = excerpt_indices[0]
    group_end   = excerpt_indices[0]

    for idx in excerpt_indices[1:]:
        # All paragraphs between group_end and idx must be empty for the group to continue
        gap_is_empty = all(
            not paragraphs[k].text.strip()
            for k in range(group_end + 1, idx)
        )
        if gap_is_empty:
            group_end = idx
        else:
            groups.append((group_start, group_end))
            group_start = idx
            group_end   = idx

    groups.append((group_start, group_end))

    # Apply indent to every paragraph within each group (excerpt and empty gaps)
    count = 0
    for group_start, group_end in groups:
        for i in range(group_start, group_end + 1):
            p = paragraphs[i]
            p.paragraph_format.left_indent  = _EXCERPT_LEFT_INDENT
            p.paragraph_format.right_indent = _EXCERPT_RIGHT_INDENT
            if i in excerpt_set:
                log.debug("Para %d: applied excerpt block indent", i)
                count += 1

    # Insert blank paragraphs before/after each group where missing.
    # Process in forward order; addnext/addprevious operate on XML siblings
    # so earlier insertions don't shift later group indices.
    for group_start, group_end in groups:
        first_elem = paragraphs[group_start]._element
        last_elem  = paragraphs[group_end]._element

        prev = first_elem.getprevious()
        if prev is None or not _elem_is_empty_para(prev):
            first_elem.addprevious(OxmlElement("w:p"))
            log.debug("Para %d: inserted blank paragraph before excerpt group", group_start)

        nxt = last_elem.getnext()
        if nxt is None or not _elem_is_empty_para(nxt):
            last_elem.addnext(OxmlElement("w:p"))
            log.debug("Para %d: inserted blank paragraph after excerpt group", group_end)

    return count


# ---------------------------------------------------------------------------
# Feature: font conversion
# ---------------------------------------------------------------------------

def convert_fonts(doc, target_font, skip_fonts):
    """Replace all fonts in the document with target_font, except those listed
    in skip_fonts (case-insensitive).

    Covers all four locations where fonts live in a docx:
      1. Document defaults  (w:docDefaults in the styles part)
      2. Style definitions  (w:style elements)
      3. Paragraph-level run properties  (w:pPr/w:rPr)
      4. Individual run properties  (w:r/w:rPr)

    If ANY font attribute on a w:rFonts element matches a skip font, the entire
    element is left untouched — this preserves monospace runs wholesale.

    Returns the number of w:rFonts elements updated.
    """
    skip_lower = {f.lower() for f in skip_fonts}
    FONT_ATTRS = (qn("w:ascii"), qn("w:hAnsi"), qn("w:cs"), qn("w:eastAsia"))
    count = 0

    def process(rFonts):
        nonlocal count
        for attr in FONT_ATTRS:
            val = rFonts.get(attr, "")
            if val.lower() in skip_lower:
                log.debug("Font: skipping preserved font \"%s\"", val)
                return
        changed = False
        for attr in FONT_ATTRS:
            old = rFonts.get(attr)
            if old:
                rFonts.set(attr, target_font)
                log.debug("Font: \"%s\" → \"%s\"", old, target_font)
                changed = True
        if changed:
            count += 1

    for rFonts in doc.styles.element.findall(".//" + qn("w:rFonts")):
        process(rFonts)
    for rFonts in doc.element.body.findall(".//" + qn("w:rFonts")):
        process(rFonts)

    return count


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Format and clean Word .docx files for publishing.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("-c", "--config", metavar="FILE",
                        help="YAML config file (CLI flags override config values)")
    parser.add_argument("-i", "--input",  metavar="FILE", default=None,
                        help="Input .docx file")
    parser.add_argument("-o", "--output", metavar="FILE", default=None,
                        help="Output .docx file")
    parser.add_argument("--start", metavar="TEXT", default=None,
                        help="Begin range at the first heading containing TEXT")
    parser.add_argument("--end",   metavar="TEXT", default=None,
                        help="End range before the first heading containing TEXT")
    parser.add_argument("--clean",       action=argparse.BooleanOptionalAction, default=None,
                        help="Remove extra empty paragraphs within the range")
    parser.add_argument("--page-breaks", action=argparse.BooleanOptionalAction, default=None,
                        help="Insert a page break before each Heading 1 within the range")
    parser.add_argument("--fix-hrules",  action=argparse.BooleanOptionalAction, default=None,
                        help="Replace fixed-width HR shapes with margin-relative borders")
    parser.add_argument("--indent-paragraphs", action=argparse.BooleanOptionalAction, default=None,
                        help="Add a first-line indent to body text paragraphs within the range")
    parser.add_argument("--excerpt-font", metavar="FONT", default=None,
                        help="Font that identifies embedded written excerpts; "
                             "those paragraphs get 0.5\" left/right indent with blank lines before/after")
    parser.add_argument("--font-to",   metavar="FONT", default=None,
                        help="Convert all non-skipped fonts to this font")
    parser.add_argument("--font-skip", metavar="FONT", action="append", default=None,
                        help="Font to leave unchanged (repeatable)")
    parser.add_argument("--log-level", metavar="LEVEL", default=None,
                        choices=["NONE", "ERROR", "INFO", "DEBUG"],
                        help="Logging level: NONE | ERROR | INFO | DEBUG  (default: NONE)")
    parser.add_argument("--log-file",  metavar="FILE", default=None,
                        help="Log file path (default: stderr when logging is enabled)")
    args = parser.parse_args()

    # Load config file and fill in any unset args
    if args.config:
        config = load_config(args.config)
        args = apply_config(args, config)
    else:
        for dest in ("clean", "page_breaks", "fix_hrules", "indent_paragraphs"):
            if getattr(args, dest) is None:
                setattr(args, dest, False)
        if args.font_skip is None:
            args.font_skip = []

    # Configure logging now that args are fully merged
    setup_logging(args.log_level, args.log_file)

    # Validate required fields
    if not args.input:
        parser.error("Input file required: supply -i / --input or set 'input' in config")
    if not args.output:
        parser.error("Output file required: supply -o / --output or set 'output' in config")

    if not args.clean and not args.page_breaks and not args.fix_hrules \
            and not args.indent_paragraphs and not args.excerpt_font and not args.font_to:
        parser.error("Specify at least one action: --clean, --page-breaks, --fix-hrules, "
                     "--indent-paragraphs, --excerpt-font, or --font-to")

    log.info("Input : %s", args.input)
    log.info("Output: %s", args.output)

    doc = Document(args.input)

    # Resolve range once (verbose). Re-resolve silently after mutations that
    # change paragraph count so indices stay accurate without repeating output.
    start_idx, end_idx = resolve_range(doc, args, verbose=True)

    # fix-hrules must run before clean so converted borders survive the clean pass
    if args.fix_hrules:
        count = fix_hrules(doc.paragraphs, start_idx, end_idx)
        print(f"Converted {count} horizontal line shapes to paragraph borders...")
        log.info("fix-hrules: converted %d horizontal line shapes to paragraph borders", count)

    if args.clean:
        start_idx, end_idx = resolve_range(doc, args)
        to_remove = collect_removals(doc.paragraphs, start_idx, end_idx)
        print(f"Removing {len(to_remove)} empty paragraphs...")
        log.info("clean: removing %d empty paragraphs", len(to_remove))
        for elem in to_remove:
            elem.getparent().remove(elem)

    if args.page_breaks:
        start_idx, end_idx = resolve_range(doc, args)
        count = insert_page_breaks(doc.paragraphs, start_idx, end_idx)
        print(f"Inserted {count} page breaks before Heading 1 paragraphs...")
        log.info("page-breaks: inserted %d page breaks before Heading 1 paragraphs", count)

    if args.indent_paragraphs:
        start_idx, end_idx = resolve_range(doc, args)
        count = indent_paragraphs(doc.paragraphs, start_idx, end_idx,
                                  excerpt_font=args.excerpt_font)
        print(f"Indented first line of {count} paragraphs...")
        log.info("indent-paragraphs: indented first line of %d paragraphs", count)

    if args.excerpt_font:
        start_idx, end_idx = resolve_range(doc, args)
        count = format_excerpts(doc.paragraphs, start_idx, end_idx, args.excerpt_font)
        print(f"Formatted {count} excerpt paragraphs (0.5\" block indent)...")
        log.info("excerpt-font: formatted %d excerpt paragraphs with block indent", count)

    if args.font_to:
        count = convert_fonts(doc, args.font_to, args.font_skip)
        skip_note = f"  (skipping: {', '.join(args.font_skip)})" if args.font_skip else ""
        print(f"Converted {count} font entries to '{args.font_to}'{skip_note}...")
        log.info("font: converted %d entries to '%s'%s", count, args.font_to, skip_note)

    doc.save(args.output)
    print(f"Saved : {args.output}")
    log.info("Saved: %s", args.output)


if __name__ == "__main__":
    main()
